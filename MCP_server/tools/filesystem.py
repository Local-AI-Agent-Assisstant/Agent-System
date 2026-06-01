import pytesseract
import os

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

from config import SAFE_WRITE_DIR
os.makedirs(SAFE_WRITE_DIR, exist_ok=True)

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp"
}


#------------------------------------------Write files -------------------------------------------------------------  
def next_available_name(path: str) -> str:
    """
    If file exists, generate: name_1.ext, name_2.ext, ...
    """
    if not os.path.exists(path):
        return path

    base, ext = os.path.splitext(path)
    i = 1
    while True:
        new_path = f"{base}_{i}{ext}"
        if not os.path.exists(new_path):
            return new_path
        i += 1



def write_file(filename: str, content: str = "", mode: str = "write") -> str:
    """
    One tool for text/pdf/docx.
    mode: "write" (overwrite existing file), "write_unique" (create new name if exists), or "append".
    Note: 'append' mode is NOT supported for .pdf files.
    """
    if not isinstance(content, str):
        import json
        try:
            # Try to handle a list containing a single string nicely
            if isinstance(content, list) and len(content) == 1 and isinstance(content[0], str):
                content = content[0]
            else:
                content = json.dumps(content, indent=2)
        except Exception:
            content = str(content)
            
    filename = os.path.basename(filename)
    path = os.path.join(SAFE_WRITE_DIR, filename)
    
    ext = os.path.splitext(filename)[1].lower()
    mode = (mode or "write").lower().strip()

    if mode not in {"write", "append", "write_unique"}:
        return "Error: mode must be 'write', 'write_unique', or 'append'."

    if mode == "write_unique":
        path = next_available_name(path)
        filename = os.path.basename(path)

    try:
                # ---------------- PDF ----------------
        if ext == ".pdf":
            try:
                from fpdf import FPDF
            except ImportError:
                return "Error: PDF support missing. Install: pip install fpdf2"

            if mode == "append" and os.path.exists(path):
                return "Error: PDF append needs a merge library. Use mode='write'."

            # Clean weird Windows characters (prevents black squares)
            clean_content = (content or "").encode("ascii", "ignore").decode("ascii")

            # Create PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=12)

            # multi_cell automatically wraps text and handles new pages!
            # 0 means "fill the width of the page", 7 is the line height
            pdf.multi_cell(0, 7, clean_content)

            pdf.output(path)
            return {"ok": True, "path": path, "filename": os.path.basename(path)}


        # ---------------- DOCX ----------------
        if ext == ".docx":
            try:
                from docx import Document
            except Exception:
                return "Error: DOCX support missing. Install: pip install python-docx"

            if mode == "append" and os.path.exists(path):
                doc = Document(path)   # open existing
            else:
                doc = Document()       # new file

            for line in (content or "").split("\n"):
                doc.add_paragraph(line)

            doc.save(path)
            return {"ok": True, "path": path, "filename": os.path.basename(path)}

        # ---------------- TEXT / OTHER ----------------
        py_mode = "a" if mode == "append" else "w"
        with open(path, py_mode, encoding="utf-8") as f:
            f.write(content or "")
            if content and not content.endswith("\n"):
                f.write("\n")

        return {"ok": True, "path": path, "filename": os.path.basename(path)}

    except Exception as e:
        return f"Error writing file: {e}"



# -------------------Helper for read files ------------------------------------------------------------
def pick_file_from_window(initialdir=None) -> str | None:
    from tkinter import Tk, filedialog
    try:
        root = Tk()
    except Exception:
        return None
        
    root.withdraw()
    root.attributes("-topmost", True)  # bring dialog to front
    root.update()
    file_path = filedialog.askopenfilename(
        initialdir=initialdir or os.path.expanduser("~"),
        title="Select a file to send to AI"
    )
    root.destroy()
    return file_path or None


def read_files(path: str, max_chars: int = 12000, max_pages: int = 3) -> str:
    from pypdf import PdfReader
    
    ext = os.path.splitext(path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return f"Error: Unsupported file type {ext}"

    # ---- PDF ----
    if ext == ".pdf":
        try:
            reader = PdfReader(path)
            pages = min(len(reader.pages), max_pages)
            text_parts = []
            for i in range(pages):
                page_text = reader.pages[i].extract_text() or ""
                text_parts.append(f"\n--- Page {i+1} ---\n{page_text}")
            text = "\n".join(text_parts).strip()

            if not text:
                return "Error: Could not extract text from this PDF (it may be scanned images)."

            return text[:max_chars] + ("\n\n[Truncated]" if len(text) > max_chars else "")
        except Exception as e:
            return f"Error reading PDF: {e}"

    # ---- DOCX ----
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            lines = [para.text for para in doc.paragraphs]
            text = "\n".join(lines).strip()
            if not text:
                return "Error: Could not extract text from this Word document."
            return text[:max_chars] + ("\n\n[Truncated]" if len(text) > max_chars else "")
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"
        except Exception as e:
            return f"Error reading Word file: {e}"

    # ---- IMAGES ----
    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(path)
            text = pytesseract.image_to_string(image).strip()

            if not text:
                return "Image loaded successfully, but no readable text was detected."

            return text[:max_chars] + (
                "\n\n[Truncated]" if len(text) > max_chars else ""
            )

        except ImportError:
            return "Error: Image OCR support missing. Install: pip install pillow pytesseract"

        except Exception as e:
            return f"Error reading Image: {e}"


    # ---- TEXT (fallback) ----
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            data = f.read(max_chars + 1)
        if len(data) > max_chars:
            data = data[:max_chars] + "\n\n[Truncated]"
        return data
    except Exception as e:
        return f"Error reading file: {e}"        